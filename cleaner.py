import re
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def is_pure_artifact(paragraph_text):
    """Check if the paragraph is purely a formatting artifact."""
    artifacts = [
        r'^_+$',  # Just underscores
        r'^-+$',  # Just dashes
        r'^\s*Compilation No\. \d+ Compilation date: \d+/\d+/\d+\s*$',  # Just compilation info
        r'^\s*$',  # Just whitespace
        r'^Anti-Money Laundering and Counter-Terrorism Financing Act 2006\s*\d*\s*$'  # Footer line
    ]
    
    return any(re.search(pattern, paragraph_text) for pattern in artifacts)

def clean_text(text):
    """Clean text while preserving ALL numbering and structure."""
    # Remove only extra whitespace between words
    text = re.sub(r'[ \t]+', ' ', text)
    
    # Remove trailing whitespace
    text = text.rstrip()
    
    # Preserve leading whitespace for indentation
    if text.startswith(' '):
        leading_spaces = len(text) - len(text.lstrip())
        text = ' ' * leading_spaces + text.lstrip()
    
    return text

def clean_docx(input_file, output_file):
    try:
        # Load the document
        doc = Document(input_file)
        
        # Create a new document for cleaned content
        new_doc = Document()
        
        # Track previous paragraph to handle consecutive artifacts
        prev_was_artifact = False
        
        # Process each paragraph
        for paragraph in doc.paragraphs:
            text = paragraph.text
            
            # Check if it's a pure formatting artifact
            if is_pure_artifact(text):
                prev_was_artifact = True
                continue
            
            # Clean up the text while preserving structure
            cleaned_text = clean_text(text)
            
            # If there's content after cleaning
            if cleaned_text:
                new_para = new_doc.add_paragraph()
                new_para.alignment = paragraph.alignment
                
                # Preserve paragraph style if it exists
                if paragraph.style:
                    new_para.style = paragraph.style
                
                # Copy the text with its formatting
                runs = paragraph.runs
                if runs:
                    # If there are formatted runs, preserve their formatting
                    for run in runs:
                        new_run = new_para.add_run(run.text)
                        new_run.bold = run.bold
                        new_run.italic = run.italic
                        new_run.underline = run.underline
                        
                        if run.font.name:
                            new_run.font.name = run.font.name
                        if run.font.size:
                            new_run.font.size = run.font.size
                else:
                    # If no runs, just add the cleaned text
                    new_para.add_run(cleaned_text)
                
                prev_was_artifact = False
        
        # Save the cleaned document
        new_doc.save(output_file)
        print(f"Successfully cleaned document and saved to {output_file}")
        
    except Exception as e:
        print(f"Error processing file: {str(e)}")

def main():
    try:
        # Make sure python-docx is installed
        import docx
    except ImportError:
        print("Please install python-docx first: pip install python-docx")
        return

    input_file = "ASIC2.docx"
    output_file = "asic2cleaned.docx"
    clean_docx(input_file, output_file)

if __name__ == "__main__":
    main()